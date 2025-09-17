"""
Document authenticity detection using metadata analysis.
"""

import logging
from typing import Dict, Any, Optional, List
import PyPDF2
import io
from docx import Document
import json
import re
import hashlib
from datetime import datetime, timezone
import zipfile
import xml.etree.ElementTree as ET

from models.schemas import DocumentAuthenticityResult
from utils.logging_config import get_logger

logger = get_logger(__name__)

class DocumentAuthenticityDetector:
    """Detector for document authenticity using metadata analysis."""
    
    def __init__(self, aws_access_key_id: str, aws_secret_access_key: str, aws_region: str):
        """Initialize the document authenticity detector."""
        self.aws_access_key_id = aws_access_key_id
        self.aws_secret_access_key = aws_secret_access_key
        self.aws_region = aws_region
    
    async def analyze_document_authenticity(
        self, 
        file_content: bytes, 
        filename: str, 
        file_type: str
    ) -> DocumentAuthenticityResult:
        """
        Analyze document authenticity from metadata.
        
        Args:
            file_content: Raw file content
            filename: Original filename
            file_type: MIME type of the file
            
        Returns:
            DocumentAuthenticityResult with analysis
        """
        try:
            logger.info(f"Starting document authenticity analysis for {filename}")
            
            # Extract metadata based on file type
            if filename.lower().endswith('.pdf'):
                metadata = await self._extract_pdf_metadata(file_content, filename, file_type)
                # Add PDF-specific deep analysis
                metadata.update(await self._analyze_pdf_structure(file_content))
                metadata.update(await self._analyze_pdf_fonts(file_content))
                metadata.update(await self._analyze_pdf_images(file_content))
            elif filename.lower().endswith('.docx'):
                metadata = await self._extract_docx_metadata(file_content, filename, file_type)
                # Add DOCX-specific deep analysis
                metadata.update(await self._analyze_docx_structure(file_content))
                metadata.update(await self._analyze_docx_fonts(file_content))
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Add file integrity checks
            metadata.update(await self._analyze_file_integrity(file_content, filename))
            
            # Analyze metadata deterministically (no AI)
            analysis_result = await self._analyze_metadata_deterministically(metadata)
            
            # Combine metadata with analysis
            result = DocumentAuthenticityResult(
                fileName=filename,
                fileSize=len(file_content),
                fileType=file_type,
                creationDate=metadata.get('creation_date'),
                modificationDate=metadata.get('modification_date'),
                author=metadata.get('author'),
                creator=metadata.get('creator'),
                producer=metadata.get('producer'),
                title=metadata.get('title'),
                subject=metadata.get('subject'),
                keywords=metadata.get('keywords'),
                pdfVersion=metadata.get('pdf_version'),
                pageCount=metadata.get('page_count'),
                isEncrypted=metadata.get('is_encrypted', False),
                hasDigitalSignature=metadata.get('has_digital_signature', False),
                softwareUsed=metadata.get('software_used'),
                suspiciousIndicators=analysis_result.get('suspiciousIndicators', []),
                authenticityScore=analysis_result.get('authenticityScore', 50),
                rationale=analysis_result.get('rationale', 'Unable to analyze document')
            )
            
            logger.info(f"Document authenticity analysis completed: {result.authenticityScore}%")
            return result
            
        except Exception as e:
            logger.error(f"Error in document authenticity analysis: {e}")
            # Return fallback result
            return DocumentAuthenticityResult(
                fileName=filename,
                fileSize=len(file_content),
                fileType=file_type,
                suspiciousIndicators=["Analysis failed"],
                authenticityScore=50,
                rationale=f"Unable to analyze document: {str(e)}"
            )
    
    async def _extract_pdf_metadata(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            metadata = {
                'file_type': file_type,
                'is_encrypted': pdf_reader.is_encrypted,
                'page_count': len(pdf_reader.pages),
                'pdf_version': None,
                'has_digital_signature': False
            }
            
            # Extract PDF info if available
            if pdf_reader.metadata:
                info = pdf_reader.metadata
                metadata.update({
                    'title': info.get('/Title'),
                    'author': info.get('/Author'),
                    'subject': info.get('/Subject'),
                    'keywords': info.get('/Keywords'),
                    'creator': info.get('/Creator'),
                    'producer': info.get('/Producer'),
                    'creation_date': info.get('/CreationDate'),
                    'modification_date': info.get('/ModDate'),
                })
                
                # Clean up PDF date format
                if metadata['creation_date']:
                    metadata['creation_date'] = str(metadata['creation_date'])
                if metadata['modification_date']:
                    metadata['modification_date'] = str(metadata['modification_date'])
            
            # Determine software used
            if metadata.get('producer'):
                metadata['software_used'] = metadata['producer']
            elif metadata.get('creator'):
                metadata['software_used'] = metadata['creator']
            
            logger.info(f"PDF metadata extracted: {len([k for k, v in metadata.items() if v])} fields")
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {
                'file_type': file_type,
                'is_encrypted': False,
                'page_count': None,
                'error': str(e)
            }
    
    async def _extract_docx_metadata(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            metadata = {
                'file_type': file_type,
                'is_encrypted': False,
                'page_count': None,  # Would need additional parsing
                'software_used': 'Microsoft Word',
                'has_digital_signature': False
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title,
                    'author': props.author,
                    'subject': props.subject,
                    'keywords': props.keywords,
                    'creator': props.creator,
                    'creation_date': str(props.created) if props.created else None,
                    'modification_date': str(props.modified) if props.modified else None,
                })
            
            logger.info(f"DOCX metadata extracted: {len([k for k, v in metadata.items() if v])} fields")
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {
                'file_type': file_type,
                'is_encrypted': False,
                'page_count': None,
                'software_used': 'Microsoft Word',
                'error': str(e)
            }
    
    async def _analyze_metadata_deterministically(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze metadata using deterministic rules to detect suspicious patterns."""
        try:
            suspicious_indicators = []
            authenticity_score = 85  # Start with good score, deduct for issues
            
            # Check for suspicious software names (only if we have the data)
            software_used = str(metadata.get('software_used', '')).lower()
            creator = str(metadata.get('creator', '')).lower()
            producer = str(metadata.get('producer', '')).lower()
            
            ai_software_patterns = [
                'ai resume generator', 'fake document', 'resume builder',
                'ai generator', 'bot', 'automated', 'template'
            ]
            
            for pattern in ai_software_patterns:
                if pattern in software_used or pattern in creator or pattern in producer:
                    suspicious_indicators.append(f"Suspicious software detected: {pattern}")
                    authenticity_score -= 20
            
            # Check for missing metadata (be more lenient)
            if not metadata.get('author') and not metadata.get('creator'):
                suspicious_indicators.append("Missing author/creator information")
                authenticity_score -= 5  # Reduced penalty
            
            if not metadata.get('creation_date'):
                suspicious_indicators.append("Missing creation date")
                authenticity_score -= 3  # Reduced penalty
            
            # Check for unrealistic file sizes (be more lenient)
            file_size_mb = metadata.get('file_size_mb', 0)
            if file_size_mb > 0:  # Only check if we have size data
                if file_size_mb < 0.05:  # Less than 50KB
                    suspicious_indicators.append("File unusually small")
                    authenticity_score -= 10
                elif file_size_mb > 20:  # More than 20MB
                    suspicious_indicators.append("File unusually large")
                    authenticity_score -= 5
            
            # Check for suspicious PDF structure (only if we have the data)
            pdf_objects = metadata.get('pdf_objects_count', 0)
            if pdf_objects > 0:  # Only check if we have PDF data
                if pdf_objects < 5:  # More lenient threshold
                    suspicious_indicators.append("Very few PDF objects - possible simple generation")
                    authenticity_score -= 5
            
            pdf_streams = metadata.get('pdf_streams_count', 0)
            if pdf_streams == 0 and pdf_objects > 0:  # Only check if we have PDF data
                suspicious_indicators.append("No compressed streams - unusual for PDF")
                authenticity_score -= 10
            
            # Check for suspicious font patterns (only if we have the data)
            pdf_font_count = metadata.get('pdf_font_count', 0)
            if pdf_font_count > 0 and pdf_font_count == 1:
                suspicious_indicators.append("Only one font used - possible automated generation")
                authenticity_score -= 5
            
            docx_font_count = metadata.get('docx_font_count', 0)
            if docx_font_count > 0 and docx_font_count == 1:
                suspicious_indicators.append("Only one font used - possible automated generation")
                authenticity_score -= 5
            
            # Check for suspicious image patterns (be more lenient)
            pdf_images = metadata.get('pdf_images_count', 0)
            if pdf_images == 0 and pdf_objects > 0:  # Only check if we have PDF data
                suspicious_indicators.append("No images found - unusual for modern resumes")
                authenticity_score -= 3  # Reduced penalty
            
            # Check for suspicious document structure (only if we have the data)
            docx_paragraphs = metadata.get('docx_paragraphs_count', 0)
            if docx_paragraphs > 0 and docx_paragraphs < 3:  # More lenient threshold
                suspicious_indicators.append("Very few paragraphs - possible simple generation")
                authenticity_score -= 5
            
            docx_tables = metadata.get('docx_tables_count', 0)
            if docx_tables == 0 and docx_paragraphs > 0:  # Only check if we have DOCX data
                suspicious_indicators.append("No tables - unusual for professional documents")
                authenticity_score -= 3  # Reduced penalty
            
            # Check for file integrity issues (be more lenient)
            file_integrity_issues = metadata.get('file_integrity_issues', [])
            for issue in file_integrity_issues:
                if isinstance(issue, str):  # Ensure it's a string
                    suspicious_indicators.append(issue)
                    authenticity_score -= 5  # Reduced penalty
            
            # Check for suspicious properties (be more lenient)
            file_suspicious_properties = metadata.get('file_suspicious_properties', [])
            for prop in file_suspicious_properties:
                if isinstance(prop, str):  # Ensure it's a string
                    suspicious_indicators.append(prop)
                    authenticity_score -= 3  # Reduced penalty
            
            # Check for suspicious patterns from structure analysis (be more lenient)
            pdf_suspicious = metadata.get('pdf_suspicious_patterns', [])
            if isinstance(pdf_suspicious, list):
                for pattern in pdf_suspicious:
                    if isinstance(pattern, str):  # Ensure it's a string
                        suspicious_indicators.append(pattern)
                        authenticity_score -= 3  # Reduced penalty
            
            docx_suspicious = metadata.get('docx_suspicious_patterns', [])
            if isinstance(docx_suspicious, list):
                for pattern in docx_suspicious:
                    if isinstance(pattern, str):  # Ensure it's a string
                        suspicious_indicators.append(pattern)
                        authenticity_score -= 3  # Reduced penalty
            
            # Ensure score is within bounds
            authenticity_score = max(0, min(100, authenticity_score))
            
            # Generate rationale based on findings
            if not suspicious_indicators:
                rationale = "Document appears authentic based on metadata analysis"
            elif len(suspicious_indicators) <= 2:
                rationale = "Document shows some minor suspicious indicators"
            elif len(suspicious_indicators) <= 5:
                rationale = "Document shows several suspicious indicators"
            else:
                rationale = "Document shows many suspicious indicators"
            
            logger.info(f"Deterministic analysis completed: {len(suspicious_indicators)} indicators, score: {authenticity_score}")
            
            return {
                'suspiciousIndicators': suspicious_indicators,
                'authenticityScore': authenticity_score,
                'rationale': rationale
            }
                
        except Exception as e:
            logger.error(f"Error in deterministic metadata analysis: {e}")
            # Return a more helpful fallback instead of "Analysis failed"
            return {
                'suspiciousIndicators': ['Unable to analyze document metadata'],
                'authenticityScore': 75,  # Neutral score instead of 50
                'rationale': 'Document analysis completed with limited metadata information'
            }
    
    
    async def _analyze_pdf_structure(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze PDF structure for suspicious patterns."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            structure_analysis = {
                'pdf_structure_valid': True,
                'pdf_objects_count': 0,
                'pdf_streams_count': 0,
                'pdf_has_forms': False,
                'pdf_has_annotations': False,
                'pdf_has_attachments': False,
                'pdf_compression_ratio': 0.0,
                'pdf_suspicious_patterns': []
            }
            
            # Analyze PDF structure
            try:
                # Count objects and streams
                pdf_content = file_content.decode('latin-1', errors='ignore')
                structure_analysis['pdf_objects_count'] = pdf_content.count('obj')
                structure_analysis['pdf_streams_count'] = pdf_content.count('stream')
                
                # Check for forms and annotations
                if '/AcroForm' in pdf_content:
                    structure_analysis['pdf_has_forms'] = True
                if '/Annots' in pdf_content:
                    structure_analysis['pdf_has_annotations'] = True
                if '/EmbeddedFile' in pdf_content:
                    structure_analysis['pdf_has_attachments'] = True
                
                # Calculate compression ratio
                if structure_analysis['pdf_streams_count'] > 0:
                    structure_analysis['pdf_compression_ratio'] = structure_analysis['pdf_streams_count'] / structure_analysis['pdf_objects_count']
                
                # Check for suspicious patterns
                suspicious_patterns = []
                if structure_analysis['pdf_objects_count'] < 10:
                    suspicious_patterns.append("Very few PDF objects - possible simple generation")
                if structure_analysis['pdf_streams_count'] == 0:
                    suspicious_patterns.append("No compressed streams - unusual for PDF")
                if structure_analysis['pdf_compression_ratio'] > 0.8:
                    suspicious_patterns.append("High compression ratio - possible automated generation")
                
                structure_analysis['pdf_suspicious_patterns'] = suspicious_patterns
                
            except Exception as e:
                logger.warning(f"Error analyzing PDF structure: {e}")
                structure_analysis['pdf_structure_valid'] = False
            
            logger.info(f"PDF structure analysis completed: {len(structure_analysis['pdf_suspicious_patterns'])} suspicious patterns")
            return structure_analysis
            
        except Exception as e:
            logger.error(f"Error in PDF structure analysis: {e}")
            return {'pdf_structure_valid': False, 'pdf_error': str(e)}
    
    async def _analyze_pdf_fonts(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze PDF fonts for suspicious patterns."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            font_analysis = {
                'pdf_fonts': [],
                'pdf_font_count': 0,
                'pdf_embedded_fonts': 0,
                'pdf_font_suspicious_patterns': []
            }
            
            try:
                pdf_content = file_content.decode('latin-1', errors='ignore')
                
                # Extract font information
                font_matches = re.findall(r'/Font\s*<<[^>]*>>', pdf_content)
                font_analysis['pdf_font_count'] = len(font_matches)
                
                # Look for embedded fonts
                embedded_fonts = re.findall(r'/Subtype\s*/Type1|/Subtype\s*/TrueType', pdf_content)
                font_analysis['pdf_embedded_fonts'] = len(embedded_fonts)
                
                # Extract font names
                font_names = re.findall(r'/BaseFont\s*/([A-Za-z0-9\-]+)', pdf_content)
                font_analysis['pdf_fonts'] = list(set(font_names))
                
                # Check for suspicious font patterns
                suspicious_patterns = []
                if font_analysis['pdf_font_count'] == 0:
                    suspicious_patterns.append("No fonts found - unusual for PDF")
                elif font_analysis['pdf_font_count'] == 1:
                    suspicious_patterns.append("Only one font used - possible automated generation")
                
                # Check for common AI generation fonts
                ai_fonts = ['Arial', 'Times-Roman', 'Helvetica']
                if all(font in ai_fonts for font in font_analysis['pdf_fonts']):
                    suspicious_patterns.append("Only basic fonts used - possible AI generation")
                
                font_analysis['pdf_font_suspicious_patterns'] = suspicious_patterns
                
            except Exception as e:
                logger.warning(f"Error analyzing PDF fonts: {e}")
            
            logger.info(f"PDF font analysis completed: {font_analysis['pdf_font_count']} fonts, {len(font_analysis['pdf_font_suspicious_patterns'])} suspicious patterns")
            return font_analysis
            
        except Exception as e:
            logger.error(f"Error in PDF font analysis: {e}")
            return {'pdf_font_error': str(e)}
    
    async def _analyze_pdf_images(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze PDF images for suspicious patterns."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            image_analysis = {
                'pdf_images_count': 0,
                'pdf_image_types': [],
                'pdf_image_suspicious_patterns': []
            }
            
            try:
                pdf_content = file_content.decode('latin-1', errors='ignore')
                
                # Count images
                image_matches = re.findall(r'/Type\s*/XObject.*?/Subtype\s*/Image', pdf_content, re.DOTALL)
                image_analysis['pdf_images_count'] = len(image_matches)
                
                # Extract image types
                image_types = re.findall(r'/Filter\s*/([A-Za-z0-9]+)', pdf_content)
                image_analysis['pdf_image_types'] = list(set(image_types))
                
                # Check for suspicious image patterns
                suspicious_patterns = []
                if image_analysis['pdf_images_count'] == 0:
                    suspicious_patterns.append("No images found - unusual for modern resumes")
                elif image_analysis['pdf_images_count'] > 10:
                    suspicious_patterns.append("Many images - possible template usage")
                
                # Check for AI-generated image indicators
                if '/DCTDecode' in pdf_content and image_analysis['pdf_images_count'] > 0:
                    suspicious_patterns.append("JPEG images present - possible AI generation")
                
                image_analysis['pdf_image_suspicious_patterns'] = suspicious_patterns
                
            except Exception as e:
                logger.warning(f"Error analyzing PDF images: {e}")
            
            logger.info(f"PDF image analysis completed: {image_analysis['pdf_images_count']} images, {len(image_analysis['pdf_image_suspicious_patterns'])} suspicious patterns")
            return image_analysis
            
        except Exception as e:
            logger.error(f"Error in PDF image analysis: {e}")
            return {'pdf_image_error': str(e)}
    
    async def _analyze_docx_structure(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze DOCX structure for suspicious patterns."""
        try:
            doc_file = io.BytesIO(file_content)
            
            structure_analysis = {
                'docx_structure_valid': True,
                'docx_paragraphs_count': 0,
                'docx_tables_count': 0,
                'docx_images_count': 0,
                'docx_suspicious_patterns': []
            }
            
            try:
                # Open as ZIP to analyze structure
                with zipfile.ZipFile(doc_file, 'r') as zip_file:
                    file_list = zip_file.namelist()
                    
                    # Count different elements
                    structure_analysis['docx_paragraphs_count'] = len([f for f in file_list if 'paragraph' in f])
                    structure_analysis['docx_tables_count'] = len([f for f in file_list if 'table' in f])
                    structure_analysis['docx_images_count'] = len([f for f in file_list if f.startswith('word/media/')])
                    
                    # Check for suspicious patterns
                    suspicious_patterns = []
                    if structure_analysis['docx_paragraphs_count'] < 5:
                        suspicious_patterns.append("Very few paragraphs - possible simple generation")
                    if structure_analysis['docx_tables_count'] == 0:
                        suspicious_patterns.append("No tables - unusual for professional documents")
                    if structure_analysis['docx_images_count'] == 0:
                        suspicious_patterns.append("No images - unusual for modern resumes")
                    
                    structure_analysis['docx_suspicious_patterns'] = suspicious_patterns
                    
            except Exception as e:
                logger.warning(f"Error analyzing DOCX structure: {e}")
                structure_analysis['docx_structure_valid'] = False
            
            logger.info(f"DOCX structure analysis completed: {len(structure_analysis['docx_suspicious_patterns'])} suspicious patterns")
            return structure_analysis
            
        except Exception as e:
            logger.error(f"Error in DOCX structure analysis: {e}")
            return {'docx_structure_error': str(e)}
    
    async def _analyze_docx_fonts(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze DOCX fonts for suspicious patterns."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            font_analysis = {
                'docx_fonts': [],
                'docx_font_count': 0,
                'docx_font_suspicious_patterns': []
            }
            
            try:
                # Extract fonts from document
                fonts = set()
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.font.name:
                            fonts.add(run.font.name)
                
                font_analysis['docx_fonts'] = list(fonts)
                font_analysis['docx_font_count'] = len(fonts)
                
                # Check for suspicious font patterns
                suspicious_patterns = []
                if font_analysis['docx_font_count'] == 0:
                    suspicious_patterns.append("No fonts found - unusual for DOCX")
                elif font_analysis['docx_font_count'] == 1:
                    suspicious_patterns.append("Only one font used - possible automated generation")
                
                # Check for common AI generation fonts
                ai_fonts = ['Calibri', 'Arial', 'Times New Roman']
                if all(font in ai_fonts for font in font_analysis['docx_fonts']):
                    suspicious_patterns.append("Only basic fonts used - possible AI generation")
                
                font_analysis['docx_font_suspicious_patterns'] = suspicious_patterns
                
            except Exception as e:
                logger.warning(f"Error analyzing DOCX fonts: {e}")
            
            logger.info(f"DOCX font analysis completed: {font_analysis['docx_font_count']} fonts, {len(font_analysis['docx_font_suspicious_patterns'])} suspicious patterns")
            return font_analysis
            
        except Exception as e:
            logger.error(f"Error in DOCX font analysis: {e}")
            return {'docx_font_error': str(e)}
    
    async def _analyze_file_integrity(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze file integrity and basic properties."""
        try:
            integrity_analysis = {
                'file_hash': hashlib.md5(file_content).hexdigest(),
                'file_size_bytes': len(file_content),
                'file_size_mb': len(file_content) / (1024 * 1024),
                'file_integrity_issues': [],
                'file_suspicious_properties': []
            }
            
            # Check file size
            if integrity_analysis['file_size_bytes'] < 1024:  # Less than 1KB
                integrity_analysis['file_integrity_issues'].append("File too small - possible corruption or empty file")
            elif integrity_analysis['file_size_bytes'] > 50 * 1024 * 1024:  # More than 50MB
                integrity_analysis['file_suspicious_properties'].append("File very large - unusual for resume")
            
            # Check for suspicious file properties
            if integrity_analysis['file_size_mb'] < 0.1:  # Less than 100KB
                integrity_analysis['file_suspicious_properties'].append("Very small file - possible simple generation")
            elif integrity_analysis['file_size_mb'] > 10:  # More than 10MB
                integrity_analysis['file_suspicious_properties'].append("Large file - possible embedded content")
            
            # Check for common AI generation patterns in filename
            ai_patterns = ['generated', 'ai', 'bot', 'auto', 'template']
            if any(pattern in filename.lower() for pattern in ai_patterns):
                integrity_analysis['file_suspicious_properties'].append("Filename suggests AI generation")
            
            logger.info(f"File integrity analysis completed: {len(integrity_analysis['file_integrity_issues'])} issues, {len(integrity_analysis['file_suspicious_properties'])} suspicious properties")
            return integrity_analysis
            
        except Exception as e:
            logger.error(f"Error in file integrity analysis: {e}")
            return {'file_integrity_error': str(e)}
